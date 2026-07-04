"""
report_generator.py
===================
VIP Multi-Format Report Generator
Generates reports in HTML, PDF, Excel, Word, JSON, and CSV formats

Author: [Your Name]
Date: December 2024
"""

from datetime import datetime
import json
import csv
import sqlite3
from pathlib import Path
from config import REPORT_CONFIG

# PDF Generation
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib import colors
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# DOCX Generation
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Excel Generation
try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False


class VIPReportGenerator:
    """Generate security reports in multiple formats"""
    
    def __init__(self, db_path='scan_results.db'):
        """
        Initialize report generator
        
        Args:
            db_path (str): Path to SQLite database
        """
        self.db_path = db_path
        self.output_dir = REPORT_CONFIG['output_dir']
        
        # Create output directory if it doesn't exist
        Path(self.output_dir).mkdir(parents=True, exist_ok=True)
        
        print("[+] Report Generator initialized")
    
    def get_scan_data(self, scan_id):
        """
        Retrieve scan data from database
        
        Args:
            scan_id (int): Scan ID
            
        Returns:
            dict: Scan data with vulnerabilities
        """
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        # Get scan info
        cursor.execute('SELECT * FROM scans WHERE id = ?', (scan_id,))
        scan = cursor.fetchone()
        
        if not scan:
            conn.close()
            return None
        
        # Get vulnerabilities
        cursor.execute('SELECT * FROM vulnerabilities WHERE scan_id = ?', (scan_id,))
        vulns = cursor.fetchall()
        
        data = {
            'scan_id': scan[0],
            'target_url': scan[1],
            'scan_type': scan[2],
            'start_time': scan[3],
            'end_time': scan[4],
            'total_alerts': scan[5],
            'high_risk': scan[6],
            'medium_risk': scan[7],
            'low_risk': scan[8],
            'status': scan[9],
            'vulnerabilities': []
        }
        
        for v in vulns:
            data['vulnerabilities'].append({
                'id': v[0],
                'name': v[2],
                'severity': v[3],
                'confidence': v[4],
                'url': v[5],
                'description': v[6],
                'solution': v[7],
                'reference': v[8]
            })
        
        conn.close()
        return data
    
    def generate_html_report(self, scan_id, output_file=None):
        """Generate VIP HTML report with 3D styling"""
        if not output_file:
            output_file = f'{self.output_dir}report_{scan_id}.html'
        
        data = self.get_scan_data(scan_id)
        if not data:
            print(f"[!] Scan {scan_id} not found")
            return False
        
        html_content = self._create_html_template(data, scan_id)
        
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        print(f"[+] HTML Report: {output_file}")
        return True
    
    def _create_html_template(self, data, scan_id):
        """Create HTML template with VIP design"""
        
        # Build vulnerabilities HTML
        vulns_html = ''
        for i, v in enumerate(data['vulnerabilities'], 1):
            severity_lower = v['severity'].lower()
            
            solution_html = ''
            if v['solution']:
                solution_html = f'''
                <div style="background: #d4fc79; padding: 15px; border-radius: 10px; margin-top: 15px;">
                    <strong>Solution:</strong> {v['solution']}
                </div>'''
            
            vulns_html += f'''
            <div class="vuln-card {severity_lower}">
                <div style="display: flex; justify-content: space-between; align-items: center; margin-bottom: 20px;">
                    <h3>{i}. {v['name']}</h3>
                    <span class="severity-badge {severity_lower}">{v['severity']}</span>
                </div>
                <p><strong>Description:</strong> {v['description']}</p>
                <p><strong>Location:</strong> {v['url']}</p>
                <p><strong>Confidence:</strong> {v['confidence']}</p>
                {solution_html}
            </div>
            '''
        
        return f'''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Security Report - {data['target_url']}</title>
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{ 
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            padding: 40px 20px; min-height: 100vh;
        }}
        .container {{ 
            max-width: 1200px; margin: 0 auto; background: white;
            border-radius: 20px; box-shadow: 0 20px 60px rgba(0,0,0,0.3); overflow: hidden;
        }}
        .header {{ 
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white; padding: 60px 40px; text-align: center;
        }}
        .header h1 {{ font-size: 3em; margin-bottom: 10px; text-shadow: 2px 2px 4px rgba(0,0,0,0.3); }}
        .summary {{ 
            display: grid; grid-template-columns: repeat(auto-fit, minmax(250px, 1fr));
            gap: 30px; padding: 40px; background: #f8f9fa;
        }}
        .stat-card {{ 
            background: white; padding: 30px; border-radius: 15px;
            box-shadow: 0 10px 30px rgba(0,0,0,0.1); transition: all 0.3s ease;
        }}
        .stat-card:hover {{ transform: translateY(-10px); box-shadow: 0 15px 40px rgba(0,0,0,0.2); }}
        .stat-card .number {{ 
            font-size: 3em; font-weight: bold;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            -webkit-background-clip: text; -webkit-text-fill-color: transparent;
        }}
        .vulnerabilities {{ padding: 40px; }}
        .vuln-card {{ 
            background: white; border-radius: 15px; padding: 30px; margin-bottom: 30px;
            box-shadow: 0 5px 20px rgba(0,0,0,0.1); border-left: 6px solid #ddd;
        }}
        .vuln-card.high {{ border-left-color: #f5576c; }}
        .vuln-card.medium {{ border-left-color: #fcb69f; }}
        .vuln-card.low {{ border-left-color: #a8edea; }}
        .severity-badge {{ 
            padding: 8px 20px; border-radius: 25px; font-weight: bold;
            font-size: 0.9em; text-transform: uppercase;
        }}
        .severity-badge.high {{ background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%); color: white; }}
        .severity-badge.medium {{ background: linear-gradient(135deg, #ffecd2 0%, #fcb69f 100%); color: #8B4513; }}
        .severity-badge.low {{ background: linear-gradient(135deg, #a8edea 0%, #fed6e3 100%); color: #333; }}
        .footer {{ background: #2c3e50; color: white; padding: 40px; text-align: center; }}
        @media print {{ body {{ background: white; }} .container {{ box-shadow: none; }} }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>SECURITY SCAN REPORT</h1>
            <p>Comprehensive Vulnerability Assessment</p>
        </div>
        
        <div class="summary">
            <div class="stat-card"><h3>Total Issues</h3><div class="number">{data['total_alerts']}</div></div>
            <div class="stat-card"><h3>High Risk</h3><div class="number" style="background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%); -webkit-background-clip: text; -webkit-text-fill-color: transparent;">{data['high_risk']}</div></div>
            <div class="stat-card"><h3>Medium Risk</h3><div class="number" style="background: linear-gradient(135deg, #ffecd2 0%, #fcb69f 100%); -webkit-background-clip: text; -webkit-text-fill-color: transparent;">{data['medium_risk']}</div></div>
            <div class="stat-card"><h3>Low Risk</h3><div class="number" style="background: linear-gradient(135deg, #a8edea 0%, #fed6e3 100%); -webkit-background-clip: text; -webkit-text-fill-color: transparent;">{data['low_risk']}</div></div>
        </div>
        
        <div style="padding: 40px;">
            <h2 style="text-align: center; margin-bottom: 30px;">Scan Information</h2>
            <div style="display: grid; grid-template-columns: repeat(2, 1fr); gap: 20px;">
                <div style="padding: 20px; background: #f8f9fa; border-radius: 10px;"><strong>Target URL:</strong> {data['target_url']}</div>
                <div style="padding: 20px; background: #f8f9fa; border-radius: 10px;"><strong>Scan Type:</strong> {data['scan_type'].title()}</div>
                <div style="padding: 20px; background: #f8f9fa; border-radius: 10px;"><strong>Start Time:</strong> {data['start_time']}</div>
                <div style="padding: 20px; background: #f8f9fa; border-radius: 10px;"><strong>End Time:</strong> {data['end_time']}</div>
            </div>
        </div>
        
        <div class="vulnerabilities">
            <h2 style="text-align: center; margin-bottom: 30px;">Detailed Findings</h2>
            {vulns_html}
        </div>
        
        <div class="footer">
            <p><strong>Generated by Web Security Scanner v1.0</strong></p>
            <p>Report ID: {scan_id} | {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}</p>
        </div>
    </div>
</body>
</html>'''
    
    def generate_pdf_report(self, scan_id, output_file=None):
        """Generate PDF report"""
        if not PDF_AVAILABLE:
            print("[!] PDF generation requires: pip install reportlab")
            return False
        
        if not output_file:
            output_file = f'{self.output_dir}report_{scan_id}.pdf'
        
        data = self.get_scan_data(scan_id)
        if not data:
            return False
        
        doc = SimpleDocTemplate(output_file, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        # Title
        title = Paragraph("SECURITY SCAN REPORT", styles['Title'])
        story.append(title)
        story.append(Spacer(1, 20))
        
        # Summary table
        summary_data = [
            ['Metric', 'Count'],
            ['Total Issues', str(data['total_alerts'])],
            ['High Risk', str(data['high_risk'])],
            ['Medium Risk', str(data['medium_risk'])],
            ['Low Risk', str(data['low_risk'])]
        ]
        
        table = Table(summary_data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#667eea')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(table)
        story.append(PageBreak())
        
        # Vulnerabilities
        for i, vuln in enumerate(data['vulnerabilities'], 1):
            story.append(Paragraph(f"<b>{i}. {vuln['name']}</b> [{vuln['severity']}]", styles['Heading3']))
            story.append(Paragraph(f"Description: {vuln['description']}", styles['Normal']))
            story.append(Spacer(1, 10))
        
        doc.build(story)
        print(f"[+] PDF Report: {output_file}")
        return True
    
    def generate_json_report(self, scan_id, output_file=None):
        """Generate JSON report"""
        if not output_file:
            output_file = f'{self.output_dir}report_{scan_id}.json'
        
        data = self.get_scan_data(scan_id)
        if not data:
            return False
        
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2)
        
        print(f"[+] JSON Report: {output_file}")
        return True
    
    def generate_csv_report(self, scan_id, output_file=None):
        """Generate CSV report"""
        if not output_file:
            output_file = f'{self.output_dir}report_{scan_id}.csv'
        
        data = self.get_scan_data(scan_id)
        if not data:
            return False
        
        with open(output_file, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(['Name', 'Severity', 'URL', 'Description', 'Solution'])
            for v in data['vulnerabilities']:
                writer.writerow([v['name'], v['severity'], v['url'], v['description'], v['solution']])
        
        print(f"[+] CSV Report: {output_file}")
        return True
    
    def generate_docx_report(self, scan_id, output_file=None):
        """Generate Word document report"""
        if not DOCX_AVAILABLE:
            print("[!] DOCX generation requires: pip install python-docx")
            return False
        
        if not output_file:
            output_file = f'{self.output_dir}report_{scan_id}.docx'
        
        data = self.get_scan_data(scan_id)
        if not data:
            return False
        
        doc = Document()
        doc.add_heading('SECURITY SCAN REPORT', 0)
        doc.add_heading('Summary', 1)
        
        table = doc.add_table(rows=4, cols=2)
        table.cell(0, 0).text = 'Total Issues'
        table.cell(0, 1).text = str(data['total_alerts'])
        table.cell(1, 0).text = 'High Risk'
        table.cell(1, 1).text = str(data['high_risk'])
        table.cell(2, 0).text = 'Medium Risk'
        table.cell(2, 1).text = str(data['medium_risk'])
        table.cell(3, 0).text = 'Low Risk'
        table.cell(3, 1).text = str(data['low_risk'])
        
        doc.add_page_break()
        doc.add_heading('Findings', 1)
        
        for i, v in enumerate(data['vulnerabilities'], 1):
            doc.add_heading(f"{i}. {v['name']}", 2)
            doc.add_paragraph(f"Severity: {v['severity']}")
            doc.add_paragraph(f"Description: {v['description']}")
            doc.add_paragraph('')
        
        doc.save(output_file)
        print(f"[+] DOCX Report: {output_file}")
        return True
    
    def generate_excel_report(self, scan_id, output_file=None):
        """Generate Excel report"""
        if not EXCEL_AVAILABLE:
            print("[!] Excel generation requires: pip install openpyxl")
            return False
        
        if not output_file:
            output_file = f'{self.output_dir}report_{scan_id}.xlsx'
        
        data = self.get_scan_data(scan_id)
        if not data:
            return False
        
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = 'Summary'
        
        ws['A1'] = 'SECURITY SCAN REPORT'
        ws['A1'].font = Font(size=16, bold=True)
        ws['A3'] = 'Total Issues'
        ws['B3'] = data['total_alerts']
        ws['A4'] = 'High Risk'
        ws['B4'] = data['high_risk']
        ws['A5'] = 'Medium Risk'
        ws['B5'] = data['medium_risk']
        ws['A6'] = 'Low Risk'
        ws['B6'] = data['low_risk']
        
        ws_vulns = wb.create_sheet('Vulnerabilities')
        ws_vulns.append(['Name', 'Severity', 'URL', 'Description'])
        for v in data['vulnerabilities']:
            ws_vulns.append([v['name'], v['severity'], v['url'], v['description']])
        
        wb.save(output_file)
        print(f"[+] Excel Report: {output_file}")
        return True
    
    def generate_all_formats(self, scan_id):
        """Generate reports in all available formats"""
        print(f"\n[*] Generating all report formats for Scan ID: {scan_id}")
        
        self.generate_html_report(scan_id)
        self.generate_json_report(scan_id)
        self.generate_csv_report(scan_id)
        
        if PDF_AVAILABLE:
            self.generate_pdf_report(scan_id)
        if DOCX_AVAILABLE:
            self.generate_docx_report(scan_id)
        if EXCEL_AVAILABLE:
            self.generate_excel_report(scan_id)
        
        print("\n[+] All reports generated successfully")


# Main execution
if __name__ == "__main__":
    print("""
    ╔═══════════════════════════════════════════════════════════╗
    ║           VIP REPORT GENERATOR v2.0                      ║
    ║    Formats: HTML, PDF, JSON, CSV, DOCX, Excel           ║
    ╚═══════════════════════════════════════════════════════════╝
    """)
    
    generator = VIPReportGenerator()
    
    scan_id = input("\nEnter Scan ID: ").strip()
    
    if not scan_id.isdigit():
        print("[!] Invalid Scan ID")
        exit(1)
    
    print("\nSelect Format:")
    print("1. HTML (VIP 3D Design)")
    print("2. PDF")
    print("3. JSON")
    print("4. CSV")
    print("5. Word (DOCX)")
    print("6. Excel (XLSX)")
    print("7. ALL FORMATS")
    
    choice = input("\nYour choice (1-7): ").strip()
    
    scan_id = int(scan_id)
    
    if choice == '1':
        generator.generate_html_report(scan_id)
    elif choice == '2':
        generator.generate_pdf_report(scan_id)
    elif choice == '3':
        generator.generate_json_report(scan_id)
    elif choice == '4':
        generator.generate_csv_report(scan_id)
    elif choice == '5':
        generator.generate_docx_report(scan_id)
    elif choice == '6':
        generator.generate_excel_report(scan_id)
    elif choice == '7':
        generator.generate_all_formats(scan_id)
    else:
        print("[!] Invalid choice")